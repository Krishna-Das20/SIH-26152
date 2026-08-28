"""
Renders docs/team-brief.md into docs/team-brief.docx.

The markdown file is the single source of truth. This script exists so the Word
copy can be REGENERATED rather than hand-maintained -- two hand-edited copies of
a status document drift, and a stale brief is worse than no brief (the same
reasoning as the warning at the top of PROGRESS.md).

    python scripts/build_brief_docx.py

Requires python-docx (`pip install python-docx`). It handles exactly the
constructs team-brief.md uses -- headings, tables, blockquotes, bullets, rules,
and inline bold/italic/code/links. It is not a general markdown converter, and
does not try to be: a narrow converter that is correct on this file beats a
broad one that silently mangles it.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:  # pragma: no cover
    sys.exit("python-docx is not installed.  Run:  pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent

# Any markdown file under docs/ can be rendered. Defaults to the team brief so
# the original `python scripts/build_brief_docx.py` keeps working unchanged.
DEFAULT_SRC = ROOT / "docs" / "team-brief.md"

# ── Palette ───────────────────────────────────────────────────────────────
# Matches the styled web version: indigo accent, cool neutrals, semantic
# colours reserved for status so they never compete with the accent.
INK = RGBColor(0x15, 0x1A, 0x25)
MUTED = RGBColor(0x5A, 0x63, 0x77)
ACCENT = RGBColor(0x2E, 0x3E, 0x6E)
CODE_INK = RGBColor(0x8E, 0x39, 0x2C)

SHADE_HEADER = "E7EAF4"   # table header fill
SHADE_QUOTE = "F1F3F8"    # callout fill

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
DISPLAY_FONT = "Georgia"


# ── Low-level OOXML helpers ───────────────────────────────────────────────

def shade(cell, hex_fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def set_table_borders(table, colour: str = "DCE1EA") -> None:
    """Hairline borders on every edge -- python-docx has no API for this."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), colour)
        borders.append(el)
    tbl_pr.append(borders)


def left_bar(paragraph, colour: str = "2E3E6E") -> None:
    """A thick left rule -- the callout treatment from the web version."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), colour)
    borders.append(left)
    p_pr.append(borders)


def para_shading(paragraph, hex_fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C3CBD9")
    borders.append(bottom)
    p_pr.append(borders)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(14)


# ── Inline formatting ─────────────────────────────────────────────────────

# Order matters: code first, so ** inside backticks is not treated as bold.
INLINE = re.compile(
    r"(`[^`]+`)"
    r"|(\*\*[^*]+\*\*)"
    r"|(\[[^\]]+\]\([^)]+\))"
    r"|(\*[^*]+\*)"
)


def add_inline(paragraph, text: str, *, base_size=10.5, colour=None, bold=False, italic=False):
    """Write `text` into `paragraph`, honouring inline markdown."""
    for part in INLINE.split(text):
        if not part:
            continue
        run = None
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(base_size - 1)
            run.font.color.rgb = CODE_INK
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("[") and "](" in part:
            label, _, url = part[1:].partition("](")
            run = paragraph.add_run(label)
            run.font.color.rgb = ACCENT
            run.underline = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)

        if run.font.name is None:
            run.font.name = BODY_FONT
        if run.font.size is None:
            run.font.size = Pt(base_size)
        if run.font.color.rgb is None:
            run.font.color.rgb = colour if colour is not None else INK
        if bold:
            run.bold = True
        if italic:
            run.italic = True


# ── Block parsing ─────────────────────────────────────────────────────────

def is_table_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def emit_table(doc, rows: list[list[str]]) -> None:
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j in range(width):
            text = row[j] if j < len(row) else ""
            cell = cells[j]
            cell.paragraphs[0].text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            if i == 0:
                shade(cell, SHADE_HEADER)
                add_inline(para, text, base_size=9, colour=ACCENT, bold=True)
            else:
                add_inline(para, text, base_size=9.5)


def emit_code(doc, lines: list[str]) -> None:
    """Fenced blocks become a shaded monospace panel.

    ASCII diagrams only hold their shape in a fixed-width font, and Word will
    happily reflow them into nonsense otherwise. Each source line becomes its
    own paragraph with no spacing, so the drawing survives.
    """
    for raw in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.line_spacing = 1.0
        para_shading(p, "F4F6F9")
        run = p.add_run(raw if raw else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = INK


def emit_quote(doc, lines: list[str]) -> None:
    """Blockquotes become shaded callouts with an accent rule.

    Soft-wrapped lines inside a quote are rejoined for the same reason as body
    paragraphs: emitting one Word paragraph per SOURCE line produces a ragged
    stack of short lines instead of a flowing callout, and splits inline spans.
    A `###` line and a bullet each start a new block; a blank quote line ends
    the current one.
    """
    blocks: list[tuple[str, list[str]]] = []

    def start(kind: str, text: str) -> None:
        blocks.append((kind, [text]))

    for raw in lines:
        text = raw.lstrip(">").strip()
        if not text:
            blocks.append(("break", []))
            continue
        if text.startswith("###"):
            start("heading", text.lstrip("#").strip())
        elif text.startswith(("- ", "* ")):
            start("bullet", text[2:].strip())
        elif blocks and blocks[-1][0] in ("body", "heading", "bullet"):
            blocks[-1][1].append(text)
        else:
            start("body", text)

    for kind, parts in blocks:
        if kind == "break" or not parts:
            continue
        text = " ".join(parts)
        heading = kind == "heading"
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.14 if kind != "bullet" else 0.34)
        p.paragraph_format.space_before = Pt(8 if heading else 2)
        p.paragraph_format.space_after = Pt(3)
        left_bar(p)
        para_shading(p, SHADE_QUOTE)
        add_inline(p, ("•  " + text) if kind == "bullet" else text,
                   base_size=10.5, bold=heading,
                   colour=ACCENT if heading else INK)


def convert(md: str) -> Document:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.left_margin = Inches(0.95)
        section.right_margin = Inches(0.95)
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)

    lines = md.split("\n")
    i = 0
    table_buf: list[list[str]] = []
    quote_buf: list[str] = []
    para_buf: list[str] = []

    def flush():
        nonlocal table_buf, quote_buf, para_buf
        if table_buf:
            emit_table(doc, table_buf)
            table_buf = []
        if quote_buf:
            emit_quote(doc, quote_buf)
            quote_buf = []
        if para_buf:
            # Markdown soft-wraps a paragraph across several source lines. They
            # must be REJOINED before inline parsing: a `**bold span**` that
            # opens on one line and closes on the next leaves an unmatched `**`
            # in each half, which then renders as literal asterisks in Word.
            emit_paragraph(doc, " ".join(para_buf))
            para_buf = []

    def emit_paragraph(target, text: str) -> None:
        # A wholly-italic paragraph is the slide-hint convention in this file.
        if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
            p = target.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            add_inline(p, text, base_size=9.5, colour=MUTED)
            return
        p = target.add_paragraph()
        add_inline(p, text)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush()
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i].rstrip())
                i += 1
            i += 1  # consume the closing fence
            emit_code(doc, block)
            continue

        if stripped.startswith("|"):
            if para_buf or quote_buf:
                flush()
            if not is_table_sep(stripped):
                table_buf.append(split_row(stripped))
            i += 1
            continue
        if table_buf:
            flush()

        if stripped.startswith(">"):
            if para_buf:
                flush()
            quote_buf.append(stripped)
            i += 1
            continue
        if quote_buf:
            flush()

        # A blank line is the only paragraph terminator in markdown.
        if not stripped:
            flush()
            i += 1
            continue

        if stripped.startswith("---"):
            flush()
            horizontal_rule(doc)
            i += 1
            continue

        if stripped.startswith("#"):
            flush()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            p = doc.add_paragraph()
            sizes = {1: 22, 2: 15, 3: 11.5}
            size = sizes.get(level, 11)
            p.paragraph_format.space_before = Pt({1: 0, 2: 20, 3: 13}.get(level, 12))
            p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 4}.get(level, 5))
            add_inline(p, text, base_size=size,
                       colour=ACCENT if level <= 2 else INK, bold=True)
            for run in p.runs:
                if level <= 2:
                    run.font.name = DISPLAY_FONT
            if level == 1:
                p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, stripped[2:].strip())
            i += 1
            continue

        para_buf.append(stripped)
        i += 1

    flush()
    return doc


def main() -> None:
    src = Path(sys.argv[1]).resolve() if len(sys.argv) > 1 else DEFAULT_SRC
    if not src.exists():
        sys.exit(f"Source not found: {src}")

    out = src.with_suffix(".docx")
    doc = convert(src.read_text(encoding="utf-8"))
    doc.save(out)
    size_kb = out.stat().st_size / 1024
    print(f"Wrote {out.relative_to(ROOT)}  ({size_kb:.0f} KB)")
    print(f"Source of truth is {src.relative_to(ROOT)} -- edit that, then re-run this.")


if __name__ == "__main__":
    main()
